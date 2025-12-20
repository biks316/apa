from django.http import HttpResponse
from django.shortcuts import render
import wikipedia
import math
import pickle
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sentence_transformers import SentenceTransformer
from .models import Embedding 
from sklearn.metrics.pairwise import cosine_similarity

def home(request):
    return render(request, 'dash.html')

def gantt(request):
    return render(request, 'gantt.html')

def welcome(request):
    return render(request, 'welcome.html')


def get_summary_chunks(topic, max_sentences=10):
    try:
        summary = wikipedia.summary(topic, sentences=max_sentences)
        sentences = summary.split('. ')
        chunks = [s.strip() + '.' for s in sentences if len(s.strip()) > 20]
        return chunks
    except wikipedia.exceptions.DisambiguationError as e:
        return [f"Topic is ambiguous. Try one of these: {', '.join(e.options[:5])}"]
    except wikipedia.exceptions.PageError:
        return ["Topic not found on Wikipedia."]
    except Exception as e:
        return [f"An error occurred: {str(e)}"]

def create_title_page(doc, topic):
    doc.add_paragraph("\n\n\n\n\n")
    title_paragraph = doc.add_paragraph(topic)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.runs[0].font.size = Pt(24)

    author_info = doc.add_paragraph("Author Name\nInstitution\nCourse\nProfessor\nDate")
    author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def add_abstract(doc, prompt):
    doc.add_heading("Abstract", level=1)
    abstract_para = doc.add_paragraph(prompt)
    abstract_para.paragraph_format.first_line_indent = Inches(0.5)
    doc.add_page_break()

def add_dynamic_sections(doc, topic):
    doc.add_heading(topic, level=1)
    content_chunks = get_summary_chunks(topic)

    section_titles = [
        "Background", "Key Concepts", "Applications", "Current Trends", "Challenges", "Future Directions"
    ]

    for i, chunk in enumerate(content_chunks):
        heading = section_titles[i % len(section_titles)]
        doc.add_heading(heading, level=2)
        para = doc.add_paragraph(chunk)
        para.paragraph_format.first_line_indent = Inches(0.5)

def add_references(doc):
    doc.add_page_break()
    doc.add_heading("References", level=1)
    ref_para = doc.add_paragraph("Author, A. A. (Year). Title of work. Publisher.")
    ref_para.paragraph_format.first_line_indent = Inches(0.5)

def generate_docx_report(topic, prompt):
    doc = Document()
    create_title_page(doc, topic)
    add_abstract(doc, prompt)
    add_dynamic_sections(doc, topic)
    add_references(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_embedding(topic, prompt):
    model = SentenceTransformer('all-MiniLM-L6-v2')
    combined_text = f"{topic}: {prompt}"
    return model.encode(combined_text)  # Returns NumPy array


def match_or_insert(input_vector, topic, prompt, threshold=0.85):
    all_embeddings = Embedding.objects.all()

    best_match_id = None
    highest_similarity = 0

    for record in all_embeddings:
        try:
            stored_vector = pickle.loads(record.embedding)
            similarity = cosine_similarity([input_vector], [stored_vector])[0][0]
            if similarity > highest_similarity:
                highest_similarity = similarity
                best_match_id = record.id
        except Exception as e:
            print(f"Skipping record ID {record.id}: {e}")
            continue

    if highest_similarity >= threshold:
        print(f"Match found with similarity: {highest_similarity}")
        return best_match_id
    else:
        new_record = Embedding.objects.create(
            topic=topic,
            prompt=prompt,
            value=similarity,
            match_with=best_match_id,
            embedding=pickle.dumps(input_vector)
        )
        print(f"New embedding inserted with ID: {new_record.id}")
        return new_record.id


def create_apa_report(request):
    if request.method == "POST":
        topic = request.POST.get("topic", "Untitled Report")
        prompt = request.POST.get("prompt", "")
        generate_ppt = request.POST.get("generate_ppt") == "on"

        # Step 1: Create vector
        embedding_vector = generate_embedding(topic, prompt)

        # Step 2: Match or insert into DB
        match_id = match_or_insert(embedding_vector, topic, prompt)
        print(f"Match or Insert ID: {match_id}")

        # Step 3: Generate docx file
        docx_buffer = generate_docx_report(topic, prompt)

        # Step 4: Return as file download
        response = HttpResponse(docx_buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{topic.replace(" ", "_")}_report.docx"'
        return response

    return HttpResponse("Invalid request method.", status=400)
